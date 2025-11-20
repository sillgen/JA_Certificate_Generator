#!/usr/bin/env python3
"""
Multi-format file reader for extracting student names.
Supports: .docx, .pdf, .csv, .txt, .xlsx, and image files
"""

import os
import re
from pathlib import Path
import mimetypes

# File format specific imports
try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

class StudentNameExtractor:
    """Extract student names from various file formats."""
    
    def __init__(self):
        self.supported_formats = {
            '.txt': self._read_txt,
            '.csv': self._read_csv,
            '.docx': self._read_docx,
            '.doc': self._read_docx,
            '.pdf': self._read_pdf,
            '.xlsx': self._read_excel,
            '.xls': self._read_excel,
            '.png': self._read_image,
            '.jpg': self._read_image,
            '.jpeg': self._read_image,
            '.tiff': self._read_image,
            '.bmp': self._read_image,
            '.gif': self._read_image
        }
    
    def detect_file_type(self, file_path):
        """Detect file type based on extension and MIME type."""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        # Try to determine MIME type as backup
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        print(f"File: {file_path.name}")
        print(f"Extension: {extension}")
        print(f"MIME type: {mime_type}")
        
        return extension
    
    def extract_names(self, file_path):
        """Extract student names from any supported file format."""
        try:
            file_type = self.detect_file_type(file_path)
            
            if file_type not in self.supported_formats:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            print(f"Processing {file_type} file...")
            raw_text = self.supported_formats[file_type](file_path)
            
            if not raw_text:
                raise ValueError("No text content found in file")
            
            names = self._parse_names_from_text(raw_text)
            
            if not names:
                raise ValueError("No student names found in file")
            
            print(f"Found {len(names)} student names")
            return names
            
        except Exception as e:
            print(f"Error extracting names from {file_path}: {str(e)}")
            return []
    
    def _read_txt(self, file_path):
        """Read text from .txt file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    def _read_csv(self, file_path):
        """Read text from .csv file."""
        if pd is None:
            # Fallback to basic text reading
            return self._read_txt(file_path)
        
        try:
            # Try to read CSV with pandas
            df = pd.read_csv(file_path)
            
            # Look for columns that might contain names
            name_columns = []
            for col in df.columns:
                col_lower = col.lower()
                if any(keyword in col_lower for keyword in ['name', 'student', 'first', 'last']):
                    name_columns.append(col)
            
            if name_columns:
                # Combine all name columns
                text_content = []
                for col in name_columns:
                    text_content.extend(df[col].dropna().astype(str).tolist())
                return '\n'.join(text_content)
            else:
                # If no obvious name columns, return all text content
                return df.to_string()
                
        except Exception as e:
            print(f"CSV parsing error: {e}, falling back to text reading")
            return self._read_txt(file_path)
    
    def _read_docx(self, file_path):
        """Read text from .docx file."""
        if Document is None:
            raise ImportError("python-docx package required for .docx files")
        
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            return '\n'.join(text_content)
            
        except Exception as e:
            print(f"DOCX reading error: {e}")
            return ""
    
    def _read_pdf(self, file_path):
        """Read text from .pdf file."""
        if pdfplumber is None:
            raise ImportError("pdfplumber package required for .pdf files")
        
        try:
            text_content = []
            
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            
            return '\n'.join(text_content)
            
        except Exception as e:
            print(f"PDF reading error: {e}")
            return ""
    
    def _read_excel(self, file_path):
        """Read text from .xlsx/.xls file."""
        if pd is None:
            raise ImportError("pandas package required for Excel files")
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text_content = []
            
            for sheet_name in excel_file.sheet_names:
                print(f"Reading sheet: {sheet_name}")
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Look for name columns first
                name_columns = []
                for col in df.columns:
                    col_lower = str(col).lower()
                    if any(keyword in col_lower for keyword in ['name', 'student', 'first', 'last']):
                        name_columns.append(col)
                
                if name_columns:
                    for col in name_columns:
                        text_content.extend(df[col].dropna().astype(str).tolist())
                else:
                    # Include all non-null text content
                    for col in df.columns:
                        text_content.extend(df[col].dropna().astype(str).tolist())
            
            return '\n'.join(text_content)
            
        except Exception as e:
            print(f"Excel reading error: {e}")
            return ""
    
    def _read_image(self, file_path):
        """Read text from image using OCR."""
        if Image is None or pytesseract is None:
            raise ImportError("Pillow and pytesseract packages required for image files")
        
        try:
            # Check if tesseract is available
            try:
                pytesseract.get_tesseract_version()
            except Exception:
                print("Warning: Tesseract OCR not found. Install tesseract-ocr system package.")
                return ""
            
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Perform OCR
            text = pytesseract.image_to_string(image)
            return text
            
        except Exception as e:
            print(f"Image OCR error: {e}")
            return ""
    
    def _parse_names_from_text(self, text):
        """Parse student names from extracted text."""
        names = []
        lines = text.split('\n')
        
        # Patterns that typically indicate names
        name_patterns = [
            # Full names (First Last, First Middle Last, etc.)
            r'^[A-Z][a-z]+ [A-Z][a-z]+(?:\s+[A-Z][a-z]+)*$',
            # Names with special characters (hyphens, apostrophes)
            r'^[A-Z][a-z\-\']+\s+[A-Z][a-z\-\']+(?:\s+[A-Z][a-z\-\']+)*$',
            # Names with accented characters
            r'^[A-Z][a-zÀ-ÿ\-\']+\s+[A-Z][a-zÀ-ÿ\-\']+(?:\s+[A-Z][a-zÀ-ÿ\-\']+)*$',
            # Names with mixed case (like SaMiyah)
            r'^[A-Z][a-zA-Z\-\']+\s+[A-Z][a-zA-Z\-\']+(?:\s+[A-Z][a-zA-Z\-\']+)*$'
        ]
        
        for line in lines:
            line = line.strip()
            
            # Skip empty lines and obvious non-names
            if not line or len(line) < 3:
                continue
            
            # Skip lines that look like headers or labels
            if any(keyword in line.lower() for keyword in [
                'student', 'name', 'class', 'grade', 'school', 'teacher',
                'kindergarten', 'classroom', 'list', '2025', '2026'
            ]):
                continue
            
            # Skip lines with numbers (unless they're part of a name)
            if re.search(r'^\d+', line) or re.search(r'\d{4}', line):
                continue
            
            # Check if line matches name patterns
            for pattern in name_patterns:
                if re.match(pattern, line):
                    # Additional validation: must have at least first and last name
                    parts = line.split()
                    if len(parts) >= 2 and all(len(part) >= 2 for part in parts):
                        names.append(line)
                        break
        
        # Remove duplicates while preserving order
        unique_names = []
        seen = set()
        for name in names:
            if name.lower() not in seen:
                unique_names.append(name)
                seen.add(name.lower())
        
        return unique_names
    
    def list_supported_formats(self):
        """List all supported file formats."""
        return list(self.supported_formats.keys())

def main():
    """Test the extractor with various file types."""
    extractor = StudentNameExtractor()
    
    print("Student Name Extractor")
    print("=" * 50)
    print("Supported formats:", ", ".join(extractor.list_supported_formats()))
    print()
    
    # Test with the existing class list
    test_file = "25-26 Class list.docx"
    if os.path.exists(test_file):
        print(f"Testing with: {test_file}")
        names = extractor.extract_names(test_file)
        
        if names:
            print(f"\nExtracted {len(names)} names:")
            for i, name in enumerate(names, 1):
                print(f"{i:2d}. {name}")
        else:
            print("No names found!")
    else:
        print(f"Test file {test_file} not found!")

if __name__ == "__main__":
    main()